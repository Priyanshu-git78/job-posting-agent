from config import get_llm
from docx import Document
from docx.shared import RGBColor, Pt
import re
from .starting import State
from pydantic import BaseModel, Field
from langchain.messages import  HumanMessage, SystemMessage
import copy
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import subprocess
import os


llm = get_llm()



class Bio(BaseModel):
    bio:str = Field(description="""Write a concise, keyword-rich "Bio" (professional summary) section for a resume.

    Guidelines:
    - Length: 3-5 sentences (60-100 words), sing

    def insert_skills_with_colored_headings(paragraph, placeholder, skills_text, heading_color=(0x1F, 0x4E, 0x79)):
        full_text = ''.join(r.text for r in ple paragraph
    - Tailor the content to align the candidate's actual expebiorience with the target job's required skills and responsibilities
    - Naturally weave in keywords from the job requirements (for ATS matching) — do NOT force keywords the candidate has no experience with
    - Lead with role/title + years of experience, then key technical skills, then a differentiator (certification, domain expertise, education, etc.)
    - Use active, achievement-oriented language; avoid generic filler like "hardworking team player"
    - Do not fabricate skills, tools, or experience not present in the candidate's details
    - Output ONLY the bio text — no headers, labels, or explanations""")

    bio_highlight: list[str] = Field(
        description=(
            "Verbatim phrases extracted directly from the candidate's bio that an HR "
            "reviewer should see at a glance — e.g. job titles, years of experience, "
            "key skills, certifications, quantifiable achievements, education, and "
            "notable companies. Each item must be an exact substring of the original "
            "bio text, not a paraphrase or summary."
        )
    )

class Skills(BaseModel):
    technical_skills: str = Field(
        description=(
            "A categorized, ATS-optimized technical skills block for the resume, "
            "tailored and sorted by relevance to requirement_title.\n\n"
            "FORMAT: One category per line: 'Category Name: skill1, skill2, skill3'. "
            "Separate categories with a newline (\\n). Category names are short "
            "(2-4 words), industry-standard (e.g. 'Generative AI & LLMs', "
            "'MLOps & DevOps').\n\n"
            "SOURCING RULES (priority order):\n"
            "1. Include every skill from matched_skills.\n"
            "2. Include additional technical skills explicitly present in the resume "
            "text, even if not in matched_skills.\n"
            "3. Include a non_matched_skill ONLY if the resume gives verifiable "
            "evidence it was actually used (named in a project/role/bullet). Never "
            "add it just because it appears in the target skills list — that is "
            "fabrication. If unverified, omit it.\n"
            "4. Never invent a skill absent from the resume.\n\n"
            "DEDUPLICATION (no doubling):\n"
            "- Each distinct skill appears EXACTLY ONCE across the entire output, in "
            "its single most relevant category — never repeated in a second category "
            "even if it plausibly fits both (e.g. 'Python' goes under 'Data "
            "Engineering & Cloud' OR 'Machine Learning', not both — pick the category "
            "closest to what requirement_title emphasizes).\n"
            "- Merge near-duplicate/synonym mentions from the resume into one canonical "
            "term (e.g. 'PySpark' and 'Apache Spark' → keep as a single distinct entry "
            "using the standard name; don't list both if they refer to the same tool).\n"
            "- If the resume repeats a skill across multiple sections/roles, list it "
            "only once.\n\n"
            "RELEVANCE-BASED SORTING (applies at both levels):\n"
            "- CATEGORY ORDER: categories most relevant to requirement_title appear "
            "first. A category counts as more relevant if it contains more "
            "matched_skills and/or skills that directly appear in requirement_title's "
            "skill list.\n"
            "- SKILL ORDER WITHIN A CATEGORY: within each category, order skills by "
            "relevance to requirement_title — matched_skills first, then other "
            "resume-verified skills, then any verified non_matched_skills last. Ties "
            "break by how prominently/frequently the skill appears in the resume.\n"
            "- Do not sort alphabetically — relevance order always wins.\n\n"
            "TERMINOLOGY: Use standard capitalization/spelling as seen in job "
            "postings (e.g. 'PostgreSQL', not 'postgres') to maximize keyword match.\n\n"
            "SCOPE: Technical tools, languages, frameworks, platforms, methodologies "
            "only — no soft skills, certifications, titles, or company names.\n\n"
            "LENGTH: ~5-9 categories, ~4-10 skills each. Prioritize relevance to "
            "requirement_title over dumping every tool the resume mentions."
        )
    )


class RAGExperienceBullet(BaseModel):
    action_verb: str = Field(
        description="Strong past-tense action verb describing what the candidate did — e.g. 'Built', 'Designed', 'Architected', 'Developed'."
    )
    problem: str = Field(
        description="The specific problem or business need the RAG system was built to solve. Concrete, not generic (e.g. 'manual search across 10K docs took hours' not 'improve efficiency')."
    )
    technical_approach: str = Field(
        description="The technical architecture used — e.g. vector database, embedding model, LLM, retrieval method."
    )
    cloud_platform: str = Field(
        description="Cloud provider used for deployment — e.g. AWS, GCP, Azure. Empty string if not mentioned."
    )
    deployment_method: str = Field(
        description="Deployment method/framework used — e.g. Docker, Kubernetes, CI/CD pipeline, serverless."
    )
    result: str = Field(
        description="Quantifiable outcome achieved — include a number where possible (e.g. 'reduced query time by 60%')."
    )
    highlight_phrases: list[str] = Field(
        description=(
            "Verbatim phrases extracted directly from the above fields that an HR/technical "
            "reviewer should see highlighted at a glance — key tech names, cloud platform, "
            "and the quantifiable result. Each item must be an exact substring of the field "
            "text it came from, not a paraphrase."
        )
    )

    def to_bullet(self) -> str:
        return (
            f"{self.action_verb} a Retrieval-Augmented Generation (RAG) system to solve "
            f"{self.problem}, by {self.technical_approach}. Deployed on {self.cloud_platform} "
            f"using {self.deployment_method}, achieving {self.result}."
        )
    



class taloried_recommendations(BaseModel):
    skills: str = Field(description="List the candidate's skills that match the company's job requirements, "
            "ordered from strongest match to weakest, based on how closely each skill "
            "aligns with the specific skills mentioned in the job requirements.")
    project: str = Field(description="List the candidate's projects that best demonstrate the required skills "
            "and responsibilities, ordered from most relevant to least relevant, based "
            "on the job requirements.")
    all_gaps: str = Field(description="List all gaps between the candidate's resume and the job requirements — "
            "including missing skills, missing or insufficient project experience, and "
            "missing years of relevant experience.")

class rank_summary(BaseModel):
    rank: int = Field(description="A numeric score between 0 and 1 representing the estimated likelihood of the "
            "candidate being selected, based on how well their skills, projects, and experience "
            "match the company's requirements. 0 = no chance, 1 = perfect match.")
    summary:str =Field(description="A concise, actionable summary explaining what the candidate should add, improve, "
            "or emphasize in their resume to maximize their chances of selection, based on the "
            "identified gaps and matching strengths.")



def resume_evaluator(state:State):
    doc = Document(r"supported_documents/Priyanshu_Harsana_Resume.docx")

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    text = "\n".join(full_text)
    print(text)
    system=SystemMessage(content=
        "You are a resume-building specialist. follow these steps\n"
        "1. Analyze the job requirements provided.\n"
        "2. Compare them against the candidate's resume.\n"
        "3. List matching skills and projects.\n"
        "4. List missing or weak skills/projects the candidate should improve or add."
    )
    human= HumanMessage(
        content=f"""Bio
        resume details: {text}/n
        job requirements
        company_name:{state["company_name"]},
        requirement_title:    {state["requirement_title"]},
        requirement_skills:   {state["requirement_skill"]},
        requirement_responsibilites :{state["requirement_responsibilities"]}
        """
    )
    llm_structure = llm.with_structured_output(taloried_recommendations)
    response = llm_structure.invoke([system,human])
    return {
        "resume_text": text,
        "matching_skills": response.skills,
        "matching_projects": response.project,
        "all_gaps": response.all_gaps,
        'resume_info':text
    }

def summary(state:State):
    company_name=state["company_name"]
    requirement_title=    state["requirement_title"]
    requirement_skills=  state["requirement_skill"]
    requirement_responsibilites =state["requirement_responsibilities"]
    doc = Document(r"supported_documents/Priyanshu_Harsana_Resume.docx")
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    text = "\n".join(full_text)

    skills_match= state["matching_skills"]
    project_match=state["matching_projects"]
    all_gaps=   state["all_gaps"]
    messages =[SystemMessage(content=
    "You are a resume analyzer analyst. Your job is to analyze the company's \n"
    "job requirements against the candidate's resume. Identify matching skills \n"
    "and projects, list missing or weak skills/projects, and rank and provide the summary to the candidate \n"
    "based on overall fit.\n"
    ),
    HumanMessage(content=f"""
    company requirements details:
    {requirement_title},requirement:{requirement_skills},responsbilities{requirement_responsibilites}
    candiadate["resume"]
    matching_profile details:
    skills :{skills_match}
    projects:{project_match}
    all_gaps : {all_gaps}

    """)]
    llm_structure = llm.with_structured_output(rank_summary)
    response = llm_structure.invoke(messages)
    return response

def resumeEdit(state:State):
    """edit the resume the according to the analysis and other things extracted evlaute by the LLM"""
    doc = Document(r"supported_documents/Priyanshu_Harsana_Template.docx") #[AI/ML Engineer with 2+ years of experience building and deploying production-grade Machine Learning , Generative AI and Python-based solutions across the full ML lifecycle—from ETL and feature engineering to model training, LLM applications, multimodal RAG, MLOps, deployment, and monitoring. Skilled in Python, FastAPI, PostgreSQL (pgvector), LangChain, and cloud-ready AI systems. Currently completing an MBA in Artificial Intelligence & Machine Learning, combining strong technical expertise with business strategy.]
    requirement_title=    state["requirement_title"]
    requirement_skills=  state["requirement_skill"]
    requirement_responsibilites =state["requirement_responsibilities"]
    matched_skills = state["matching_skills"]
    non_matched_skills= state["all_gaps"]
    messages = [
        SystemMessage(
            """You are an expert resume writer specializing in ATS-optimized professional bios.

    Your task: Write a concise, keyword-rich "Bio" (professional summary) section for a resume.

    Guidelines:
    - Length: 3-5 sentences (60-100 words), single paragraph
    - Tailor the content to align the candidate's actual experience with the target job's required skills and responsibilities
    - Naturally weave in keywords from the job requirements (for ATS matching) — do NOT force keywords the candidate has no experience with
    - Lead with role/title + years of experience, then key technical skills, then a differentiator (certification, domain expertise, education, etc.)
    - Use active, achievement-oriented language; avoid generic filler like "hardworking team player"
    - Do not fabricate skills, tools, or experience not present in the candidate's details
    - Output ONLY the bio text — no headers, labels, or explanations"""
        ),
        HumanMessage(
            f"""Candidate details:
    {state['resume_info']}

    Target role: {requirement_title}
    Required skills: {requirement_skills}
    Responsibilities: {requirement_responsibilites}

    Example of the tone/style/format expected (for reference only, do not copy content):
    "AI/ML Engineer with 3+ years of experience building and deploying production-grade Machine Learning, Generative AI and Python-based solutions across the full ML lifecycle—from ETL and feature engineering to model training, LLM applications, multimodal RAG, MLOps, deployment, and monitoring. Skilled in Python, FastAPI, PostgreSQL (pgvector), LangChain, and cloud-ready AI systems. Recently completed an MBA in Artificial Intelligence & Machine Learning, combining strong technical expertise with business strategy."

    Write the bio for this candidate now."""
        ),
    ]
    messages_skills = [
        SystemMessage(
            """You are an expert resume writer specializing in ATS-optimized technical skills sections.

    Your task: Produce a categorized technical skills block for a resume, tailored and sorted by relevance to the target role.

    Guidelines:

    FORMAT
    - One category per line: "Category Name: skill1, skill2, skill3"
    - Separate categories with a newline
    - Category names are short (2-4 words), industry-standard, and instantly recognizable to recruiters and ATS parsers (e.g. "Generative AI & LLMs", "MLOps & DevOps", "Cloud & Data Engineering")

    SOURCING RULES (priority order)
    1. Include every skill from matched_skills — these are confirmed overlaps between the resume and the role.
    2. Include additional technical skills explicitly present in the candidate's resume, even if not in matched_skills, when relevant to the target role.
    3. Include a skill from non_matched_skills ONLY if the resume gives verifiable evidence the candidate actually used it (named in a specific project, role, or bullet). Never add it just because the role requires it — that is fabrication. If unverified, omit it.
    4. Never invent, assume, or infer any skill absent from the resume, regardless of how standard it is for the role.

    DEDUPLICATION
    - Each distinct skill appears exactly once across the entire output, placed in its single most relevant category — never repeated across categories even if it plausibly fits more than one.
    - Merge synonyms/near-duplicates from the resume into one canonical, standard term (e.g. "PySpark" and "Apache Spark" become a single entry).
    - If a skill is repeated across multiple sections of the resume, list it only once.

    SORTING (relevance-based, never alphabetical)
    - Category order: categories most relevant to the target role come first, based on how many matched_skills and required skills they contain.
    - Skill order within a category: matched_skills first, then other resume-verified skills, then any verified non_matched_skills last.

    TERMINOLOGY
    - Use standard capitalization/spelling as seen in job postings (e.g. "PostgreSQL", not "postgres") to maximize keyword match.

    SCOPE
    - Technical tools, languages, frameworks, platforms, and methodologies only — no soft skills, certifications, job titles, or company names.

    LENGTH
    - Roughly 5-9 categories, 4-10 skills per category. Prioritize relevance to the target role over listing every tool mentioned in the resume.

    Do not fabricate skills, tools, or experience not present in the candidate's details.
    Return the result via the Skills schema only — no headers, labels, or explanations outside the structured output."""
        ),
        HumanMessage(
            f"""Candidate details:
    {state['resume_info']}

    Target role: {requirement_title}
    Required skills for this role: {requirement_skills}
    Matched skills (already confirmed present in candidate's resume): {matched_skills}
    Non-matched skills (required by role, not yet confirmed in resume — include only if resume provides evidence): {non_matched_skills}

    Example of the format expected (for reference only, do not copy content):
    "Generative AI & LLMs: LLMs, RAG, LangChain, LangGraph, Hugging Face Transformers, Prompt Engineering, Fine-Tuning
    Data Engineering & Cloud: Python, SQL, Apache Spark (PySpark), Pandas, NumPy, ETL/ELT Pipelines, PostgreSQL, BigQuery, AWS SageMaker
    MLOps & DevOps: MLflow, Docker, Git, GitHub Actions, CI/CD, Model/Prompt Versioning"

    Write the skills section for this candidate now."""
        ),
    ]
    messages_rag_experience = [
    SystemMessage(
        """You are an expert resume writer specializing in ATS-optimized experience bullets for AI/ML engineers.

    Your task: Extract and structure the candidate's Retrieval-Augmented Generation (RAG) project experience into a single, quantified experience bullet following this exact narrative arc:

    [Action verb] a Retrieval-Augmented Generation (RAG) system to solve [problem/business need], 
    by [technical approach/architecture — e.g. vector DB, embedding model, LLM]. 
    Deployed on [cloud platform] using [method/framework], achieving [quantifiable result].

    Guidelines:
    - action_verb: a strong past-tense verb (e.g. "Built", "Designed", "Architected", "Developed") that matches what the candidate actually did per their resume.
    - problem: the specific, concrete problem or business need the RAG system addressed — not generic phrasing.
    - technical_approach: the actual architecture/stack used — vector database, embedding model, LLM, retrieval method, orchestration framework — as stated in the resume.
    - cloud_platform: the cloud provider used, if mentioned. Empty string if not mentioned in the resume.
    - deployment_method: the deployment tooling/method used (e.g. Docker, Kubernetes, CI/CD), if mentioned. Empty string if not mentioned.
    - result: a quantifiable outcome (numbers, percentages, time/cost savings). If no number exists in the resume, state the qualitative outcome as concisely as possible — do not invent a metric.
    - highlight_phrases: verbatim phrases from the fields above that should be visually highlighted for HR — key tech names, cloud platform, and the result. Must be exact substrings of the field text, not paraphrased.

    Do not fabricate any detail — problem, technical approach, platform, method, or result — that is not present in the candidate's resume. If a field is not mentioned, leave it as an empty string rather than guessing.

    Tailor relevance to the target role's required skills and responsibilities where the resume supports it, but never invent alignment that isn't there.

    Return the result via the RAGExperienceBullet schema only — no headers, labels, or explanations outside the structured output."""
        ),
        HumanMessage(
            f"""Candidate details:
    {state['resume_info']}

    Target role: {requirement_title}
    Required skills: {requirement_skills}
    Responsibilities: {requirement_responsibilites}

    Extract the candidate's RAG project experience from the details above and structure it now."""
        ),
    ]


    llm_structure_bio = llm.with_structured_output(Bio)
    data = llm_structure_bio.invoke(messages)
    print('data: ', data.bio)
    llm_structure_skills = llm.with_structured_output(Skills)
    skills_obj = llm_structure_skills.invoke(messages_skills)
    print('skills :',skills_obj)
    llm_structure_rag_experience = llm.with_structured_output(RAGExperienceBullet)
    rag_experience_obj = llm_structure_rag_experience.invoke(messages_rag_experience)

    

    def insert_bio_with_highlights(paragraph, placeholder, bio_text, highlight_phrases):
        """ Highlighting the important parts of the bio"""
        full_text = ''.join(r.text for r in paragraph.runs)
        if placeholder not in full_text:
            return False

        new_text = full_text.replace(placeholder, bio_text)
        template_run = paragraph.runs[0]

        for run in paragraph.runs:
            run.text = ''

        pattern = re.compile('(' + '|'.join(re.escape(p) for p in highlight_phrases) + ')', re.IGNORECASE)
        parts = [p for p in pattern.split(new_text) if p]
        lowered_phrases = [p.lower() for p in highlight_phrases]

        first = True
        for part in parts:
            if first:
                run = paragraph.runs[0]
                run.text = part
                first = False
            else:
                run = paragraph.add_run(part)
                run.font.name = template_run.font.name
                run.font.size = template_run.font.size

            run.font.bold = part.lower() in lowered_phrases
            # If you want a highlighter-style background instead of bold:
            # from docx.enum.text import WD_COLOR_INDEX
            # run.font.highlight_color = WD_COLOR_INDEX.YELLOW if part.lower() in lowered_phrases else None

        return True
    highlight_phrases = data.bio_highlight

    llm.invoke(f" ")
    

    for paragraph in doc.paragraphs:
        insert_bio_with_highlights(paragraph, '[bio]', data.bio, highlight_phrases)

    def insert_skills_with_colored_headings(paragraph, placeholder, skills_text, heading_color=(0x1F, 0x4E, 0x79)):
        full_text = ''.join(r.text for r in paragraph.runs)
        if placeholder not in full_text:
            return False

        lines = [l.strip() for l in skills_text.strip().split('\n') if l.strip()]
        template_p = paragraph._p

        for run in paragraph.runs:
            run.text = ''

        insert_after = template_p
        for i, line in enumerate(lines):
            if ':' in line:
                heading, rest = line.split(':', 1)
                heading, rest = heading.strip() + ':', rest.strip()
            else:
                heading, rest = line, ''

            if i == 0:
                target = paragraph
            else:
                new_p = copy.deepcopy(template_p)
                for r in new_p.findall(qn('w:r')):
                    new_p.remove(r)
                insert_after.addnext(new_p)
                insert_after = new_p
                target = Paragraph(new_p, paragraph._parent)

            r1 = target.add_run(heading + ' ')
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(*heading_color)

            if rest:
                r2 = target.add_run(rest)
                r2.font.bold = False
                r2.font.color.rgb = RGBColor(0, 0, 0)

        return True


    for paragraph in doc.paragraphs:
        insert_skills_with_colored_headings(paragraph, '[skills]', skills_obj.technical_skills)



    def insert_rag_experience_with_highlights(paragraph, placeholder, rag_text, highlight_phrases):
        """Highlighting the important parts of the RAG experience bullet"""
        full_text = ''.join(r.text for r in paragraph.runs)
        if placeholder not in full_text:
            return False

        new_text = full_text.replace(placeholder, rag_text)
        template_run = paragraph.runs[0]

        for run in paragraph.runs:
            run.text = ''

        if not highlight_phrases:
            paragraph.runs[0].text = new_text
            return True

        pattern = re.compile('(' + '|'.join(re.escape(p) for p in highlight_phrases) + ')', re.IGNORECASE)
        parts = [p for p in pattern.split(new_text) if p]
        lowered_phrases = [p.lower() for p in highlight_phrases]

        first = True
        for part in parts:
            if first:
                run = paragraph.runs[0]
                run.text = part
                first = False
            else:
                run = paragraph.add_run(part)
                run.font.name = template_run.font.name
                run.font.size = template_run.font.size

            run.font.bold = part.lower() in lowered_phrases

        return True

    
    for paragraph in doc.paragraphs:
        insert_skills_with_colored_headings(paragraph, '[skills]', skills_obj.technical_skills)
        insert_rag_experience_with_highlights(paragraph, '[rag_experience]', rag_experience_obj.to_bullet(), rag_experience_obj.highlight_phrases)


    def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
        """Convert a .docx file to PDF using LibreOffice headless mode."""
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr}")

        pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Expected PDF not found at {pdf_path}")
        return pdf_path
    docx_path = "resumes/resume.docx"
    doc.save(docx_path)

    pdf_path = convert_docx_to_pdf(docx_path, "resumes")
    print("PDF saved to:", pdf_path)

    return pdf_path


